import os
import glob
import zipfile
import streamlit as st
import pandas as pd
import tempfile
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain.chains import RetrievalQA
from PyPDF2 import PdfReader
from docx import Document
from langchain.chains.question_answering import load_qa_chain
from dotenv import load_dotenv
from testscripts import script

# Initialize environment and API key
load_dotenv()

def read_pdf(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def read_txt(txt_docs):
    text = ""
    for txt in txt_docs:
        text += txt.read().decode("utf-8")  # Decode bytes to string
    return text

def read_docx(docx_docs):
    text = ""
    for docx in docx_docs:
        doc = Document(docx)
        for para in doc.paragraphs:
            text += para.text + "\n"  # Add line breaks
    return text

def read_pdf_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(directory, filename)
            pdf_reader = PdfReader(pdf_path)
            for page in pdf_reader.pages:
                text += page.extract_text()
    return text

def read_txt_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.txt'):
            txt_path = os.path.join(directory, filename)
            with open(txt_path, 'r', encoding='utf-8') as file:
                text += file.read() + "\n"  # Add line breaks
    return text

def read_docx_from_directory(directory):
    text = ""
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            docx_path = os.path.join(directory, filename)
            doc = Document(docx_path)
            for para in doc.paragraphs:
                text += para.text + "\n"  # Add line breaks
    return text

def initialize_generative_model():
    return ChatOpenAI(model="gpt-4o", temperature=0.5)

def understand_tone_and_language(script_text):
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    vector_store = FAISS.from_texts([script_text], embedding=embeddings)
    vector_store.save_local("faiss_index")

def generate_script(user_input, generative_model):
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)

    script_docs = new_db.similarity_search(user_input)

    prompt_template = """
    Take the persona of AI SAP Test Case Generator.
    A Testcase is an end to end workflow but not a single step process in SAP
    In the context PDD overviews along with their test cases are given to train the model.
    Extract the overview/purpose of the input provided.
    For the overview extracted generate the test case/s based on the above patterns of PDD overviews and Test Cases and return the text case/s as a response.
    Generate me strictly the test cases only, not test scripts or test Descriptions.
    Do not consider validations or verifications as Test Cases
    If there are any test cases already in the reference do not use them generate them based on the purpose/overview
    Also make sure that the test case does not exceed 10 words.
    The output should be in such a way that each test case should be printed in separate lines

    Example:
    Test cases for the Demo/Direct Exchange Order process:
    Test Case 1: Create a Demo Order from Customer Request via Email 
    Test Case 2: Initiate Direct Exchange Order for Customer Equipment Repair
    Test Case 3: Check for Duplicate Purchase Order Numbers in Demo/Direct Exchange Order
    Test Case 4: Determine Customer Partner Functions for Demo/Direct Exchange Order 
    Test Case 5: Process Free of Charge Equipment for Direct Exchange 
    Test Case 6: Handle Returns of Customer Equipment to Otsuka 

    Please provide at least 5 or more relevant SAP test cases, following the structure shown in the examples. Ensure to format them as "Test Case X: [Description]" for clarity.

    Context: \n{context}\n
    Generated Test Cases:
    """
    
    prompt = PromptTemplate(template=prompt_template, input_variables=["context"])
    chain = load_qa_chain(generative_model, chain_type="stuff", prompt=prompt)

    response = chain({"input_documents": script_docs}, return_only_outputs=True)
    return response["output_text"]


def test_cases_to_excel(test_cases, file_name="test_cases.xlsx"):
    # Split test cases into two columns: "Test Case Number" and "Description"
    data = []
    
    # Iterate through lines of the test cases
    for case in test_cases.splitlines():
        # Check if the line starts with "Test Case"
        if case.startswith("Test Case"):
            if ": " in case:
                # Split at the first ": " and keep the description
                description = case.split(": ", 1)[1].strip()
                test_case_number = case.split(": ", 1)[0].strip()  # Test Case X part
                data.append([test_case_number, description])  # Add test case number and description
            else:
                # Handle improperly formatted cases
                print(f"Warning: Could not split the test case line properly: {case}")
                data.append([case.strip(), ""])  # Keep the whole case and leave description blank
        else:
            # Likely a title or irrelevant line; skip or log
            print(f"Skipping non-test-case line: {case}")

    # Create a DataFrame from the list of test cases
    df = pd.DataFrame(data, columns=["Test Case Number","Description"])

    # Save the DataFrame to an Excel file
    return df





def convert_to_excel(test_scripts, file_path):
    # Create an Excel writer object with xlsxwriter engine
    with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
        # Create a summary sheet named "Test Cases" first
        summary_sheet_name = "Test Cases"
        summary_data = []
        
        # Define formats for colors
        workbook = writer.book
        header_format = workbook.add_format({'bold': True, 'bg_color': '#FFD700'})  # Gold color for headers
        title_format = workbook.add_format({'bold': True, 'bg_color': '#ADD8E6'})  # Light blue for case title

        # Loop through each test case and create a separate sheet
        for i, (case_title, script_details) in enumerate(test_scripts.items(), start=1):
            steps = []
            
            # Extract "Fiori ID", "Step Description", and "Expected Results" for each case
            for scenario, details in script_details.items():
                fiori_id = details.get("Fiori ID", "")  # Get the Fiori ID, if available
                step_desc = details.get("Step Description", {})
                exp_results = details.get("Expected Results", {})

                # Combine Fiori ID, step descriptions, and expected results in a DataFrame format
                for step_num, step_desc_text in step_desc.items():
                    steps.append({
                        
                        "Step Number": step_num,
                        "Step Description": step_desc_text,
                        "Expected Result": exp_results.get(step_num, ""),
                        "Fiori ID": fiori_id,
                    })

            # Convert data to a DataFrame
            df = pd.DataFrame(steps)
            
            # Insert the test case title at the top of the sheet
            sheet_name = f"TestScript {i}"  # Setting the sheet name to "TestScript 1", "TestScript 2", etc.
            df.to_excel(writer, index=False, sheet_name=sheet_name, startrow=1)  # Start writing data from row 2
            
            # Get the worksheet and write the test case title in the first row with color
            worksheet = writer.sheets[sheet_name]
            worksheet.write(0, 0, case_title, title_format)  # Write the test case title in the first row
            
            # Add a summary entry with "Test Case X: Title" format
            summary_data.append([f"Test Case {i}: {case_title}"])

            # Apply header format for the step columns
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(1, col_num, value, header_format)  # Apply header format to the first row of data

        # Create the summary DataFrame
        summary_df = pd.DataFrame(summary_data, columns=["Test Case Overview"])

        # Write the summary DataFrame to the summary sheet as the first sheet
        summary_df.to_excel(writer, index=False, sheet_name=summary_sheet_name, startrow=0)

        # Adjust the width of the columns in the summary sheet
        summary_worksheet = writer.sheets[summary_sheet_name]
        summary_worksheet.set_column(0, 0, 50)

        # Apply header format to the summary sheet
        summary_worksheet.write(0, 0, "Test Case Overview", header_format)


# Function definitions (read_pdf, read_txt, read_docx, etc.) go here

def main():
    st.title("Test Script Generator")

    # Choose between uploading a file or a ZIP file
    input_method = st.radio("Choose input method:", ("Upload File", "Upload ZIP File"))

    references_directory = "References"
    output_directory = "Generated Excels"

    # Ensure output directory exists
    os.makedirs(output_directory, exist_ok=True)

    # Process references
    reference_text = ""
    reference_text += read_pdf_from_directory(references_directory)
    reference_text += read_txt_from_directory(references_directory)
    reference_text += read_docx_from_directory(references_directory)

    if reference_text:
        understand_tone_and_language(reference_text)

    if input_method == "Upload File":
        uploaded_files = st.file_uploader("Upload PDF, TXT, or DOCX files", type=['pdf', 'txt', 'docx'], accept_multiple_files=True)

        if uploaded_files:
            for uploaded_file in uploaded_files:
                excel_path = process_file(uploaded_file, reference_text, output_directory)
                if excel_path:
                    with open(excel_path, 'rb') as f:
                        st.download_button(
                            label=f"Download {os.path.basename(excel_path)}",
                            data=f,
                            file_name=os.path.basename(excel_path),
                            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        )

    elif input_method == "Upload ZIP File":
        uploaded_zip = st.file_uploader("Upload a ZIP file containing your documents", type='zip')

        if uploaded_zip:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract ZIP file contents to a temporary directory
                with zipfile.ZipFile(uploaded_zip, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # Process each file in the extracted directory
                supported_extensions = ('*.pdf', '*.txt', '*.docx')
                uploaded_files = []
                for ext in supported_extensions:
                    uploaded_files.extend(glob.glob(os.path.join(temp_dir, ext)))

                if uploaded_files:
                    for uploaded_file in uploaded_files:
                        with open(uploaded_file, 'rb') as f:
                            excel_path = process_file(f, reference_text, output_directory)
                            if excel_path:
                                with open(excel_path, 'rb') as ef:
                                    st.download_button(
                                        label=f"Download {os.path.basename(excel_path)}",
                                        data=ef,
                                        file_name=os.path.basename(excel_path),
                                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                                    )
                else:
                    st.warning("No supported files found in the ZIP file.")

    st.success("All files have been processed.")

def process_file(uploaded_file, reference_text, output_directory):
    # Determine the file name and extension based on the file type
    file_name = uploaded_file.name if hasattr(uploaded_file, 'name') else os.path.basename(uploaded_file)
    st.write(f"Processing file: {file_name}")
    
    # Read content based on file extension
    uploaded_text = ""
    if file_name.endswith(".pdf"):
        uploaded_text += read_pdf([uploaded_file])
    elif file_name.endswith(".txt"):
        uploaded_text += read_txt([uploaded_file])
    elif file_name.endswith(".docx"):
        uploaded_text += read_docx([uploaded_file])
    
    if not uploaded_text:
        st.warning(f"Skipped empty or unsupported file: {file_name}")
        return None

    # Combine reference text with uploaded text
    combined_text = uploaded_text + reference_text
    understand_tone_and_language(combined_text)

    # Initialize generative model and generate test script
    generative_model = initialize_generative_model()
    response = generate_script(uploaded_text, generative_model)

    # Convert the test script to Excel
    df = test_cases_to_excel(response)
    
    testscript = {}
    for j in df["Description"]:
        testscript[j] = script(uploaded_text, j)

    # Save Excel file
    base_name = os.path.splitext(file_name)[0]
    excel_path = os.path.join(output_directory, f"{base_name}.xlsx")
    st.write(f"Saving Excel to: {excel_path}")

    convert_to_excel(testscript, excel_path)
    
    return excel_path  # Return the path of the generated Excel file

if __name__ == "__main__":
    main()
